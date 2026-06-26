from __future__ import annotations

import io
import re
import html
import difflib
from dataclasses import dataclass
from typing import Dict, Iterable, List, Tuple, Optional

from docx import Document
from pypdf import PdfReader


@dataclass
class Match:
    source_doc: str
    query_chunk: str
    source_chunk: str
    score: float

    # Campos novos para localização aproximada.
    # Eles não quebram o app atual, mas ajudam a mostrar onde está a similaridade.
    query_start_word: Optional[int] = None
    query_end_word: Optional[int] = None
    query_chunk_index: Optional[int] = None
    total_query_chunks: Optional[int] = None

    source_start_word: Optional[int] = None
    source_end_word: Optional[int] = None
    source_chunk_index: Optional[int] = None


@dataclass
class TextChunk:
    text: str
    start_word: int
    end_word: int
    chunk_index: int


def extract_text_from_txt_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text_from_docx_bytes(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_rows: List[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_rows.append(" | ".join(cells))

    return "\n".join(paragraphs + table_rows)


def extract_text_from_pdf_bytes(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(page for page in pages if page)


def normalize_text(text: str) -> str:
    text = (text or "").lower()
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def split_words(text: str) -> List[str]:
    return re.findall(r"[A-Za-zÀ-ÿ0-9]+", normalize_text(text))


def make_chunks(text: str, chunk_words: int, stride_words: int) -> List[str]:
    """
    Mantida para compatibilidade com versões anteriores do app.
    Retorna apenas os textos dos trechos.
    """
    return [chunk.text for chunk in make_positioned_chunks(text, chunk_words, stride_words)]


def make_positioned_chunks(text: str, chunk_words: int, stride_words: int) -> List[TextChunk]:
    """
    Cria trechos com localização aproximada por palavras.

    Exemplo:
    - trecho 1: palavras 1 a 60
    - trecho 2: palavras 31 a 90
    - trecho 3: palavras 61 a 120

    Isso permite ao app mostrar onde a similaridade aparece no texto analisado.
    """
    words = split_words(text)

    if not words:
        return []

    chunks: List[TextChunk] = []

    for chunk_index, start in enumerate(range(0, len(words), max(1, stride_words)), 1):
        chunk = words[start:start + chunk_words]

        if len(chunk) < max(12, chunk_words // 2):
            break

        chunks.append(
            TextChunk(
                text=" ".join(chunk),
                start_word=start + 1,
                end_word=start + len(chunk),
                chunk_index=chunk_index,
            )
        )

    # Remove duplicados preservando a primeira posição encontrada.
    unique: Dict[str, TextChunk] = {}
    for chunk in chunks:
        if chunk.text not in unique:
            unique[chunk.text] = chunk

    return list(unique.values())


def _similarity(a: str, b: str) -> float:
    return difflib.SequenceMatcher(
        None,
        normalize_text(a),
        normalize_text(b),
        autojunk=False,
    ).ratio()


def compute_matches(
    query_text: str,
    library: Dict[str, str],
    chunk_words: int,
    stride_words: int,
    top_k_per_chunk: int,
    threshold: float,
) -> Tuple[float, List[Match]]:
    """
    Compara o texto analisado com os documentos da biblioteca.

    Agora, além de retornar os trechos semelhantes, também retorna:
    - posição aproximada no texto analisado;
    - número do trecho;
    - total de trechos analisados;
    - posição aproximada no documento-fonte.
    """
    query_chunks = make_positioned_chunks(query_text, chunk_words, stride_words)

    if not query_chunks or not library:
        return 0.0, []

    source_chunks: Dict[str, List[TextChunk]] = {
        name: make_positioned_chunks(text, chunk_words, stride_words)
        for name, text in library.items()
        if text and text.strip()
    }

    matches: List[Match] = []
    matched_query_indexes = set()
    total_query_chunks = len(query_chunks)

    for query_idx, query_chunk in enumerate(query_chunks):
        candidates: List[Match] = []

        for source_name, chunks in source_chunks.items():
            for source_chunk in chunks:
                score = _similarity(query_chunk.text, source_chunk.text)

                if score >= threshold:
                    candidates.append(
                        Match(
                            source_doc=source_name,
                            query_chunk=query_chunk.text,
                            source_chunk=source_chunk.text,
                            score=score,
                            query_start_word=query_chunk.start_word,
                            query_end_word=query_chunk.end_word,
                            query_chunk_index=query_chunk.chunk_index,
                            total_query_chunks=total_query_chunks,
                            source_start_word=source_chunk.start_word,
                            source_end_word=source_chunk.end_word,
                            source_chunk_index=source_chunk.chunk_index,
                        )
                    )

        candidates.sort(key=lambda item: item.score, reverse=True)
        selected = candidates[: max(1, top_k_per_chunk)]

        if selected:
            matched_query_indexes.add(query_idx)
            matches.extend(selected)

    # Cobertura aproximada de fragmentos do texto analisado.
    # Evita somar scores sobrepostos.
    similarity = len(matched_query_indexes) / total_query_chunks

    matches.sort(key=lambda item: item.score, reverse=True)
    return similarity, matches


def get_match_location_label(match: Match) -> str:
    """
    Gera uma frase amigável para exibir no app ou no relatório.
    """
    if match.query_start_word and match.query_end_word:
        chunk_info = ""

        if match.query_chunk_index and match.total_query_chunks:
            chunk_info = f" — trecho {match.query_chunk_index} de {match.total_query_chunks}"

        return (
            f"Palavras {match.query_start_word} a {match.query_end_word}"
            f"{chunk_info}"
        )

    return "Local aproximado não identificado"


def get_source_location_label(match: Match) -> str:
    """
    Gera uma frase amigável para indicar a posição aproximada no documento-fonte.
    """
    if match.source_start_word and match.source_end_word:
        chunk_info = ""

        if match.source_chunk_index:
            chunk_info = f" — trecho {match.source_chunk_index}"

        return (
            f"Palavras {match.source_start_word} a {match.source_end_word}"
            f"{chunk_info}"
        )

    return "Local aproximado na fonte não identificado"


def highlight_text(text: str, matches: Iterable[Match]) -> str:
    """
    Destaque simples e seguro.

    Observação:
    Como os trechos comparados são normalizados, o destaque literal pode não encontrar
    exatamente a mesma forma no texto original. Por isso, esta função funciona melhor
    como apoio visual inicial, e o mapa por palavras é mais confiável.
    """
    safe = html.escape(text or "")

    for match in sorted(matches, key=lambda m: len(m.query_chunk), reverse=True):
        fragment = html.escape(match.query_chunk)

        if fragment and fragment in safe:
            safe = safe.replace(fragment, f"<mark>{fragment}</mark>")

    return safe.replace("\n", "<br>")


def summarize_matches(matches: Iterable[Match], limit: int = 10) -> List[str]:
    """
    Cria um resumo textual das correspondências para histórico, relatório ou interface.
    Não inclui o texto completo do usuário.
    """
    lines: List[str] = []

    for index, match in enumerate(list(matches)[:limit], 1):
        lines.append(
            f"{index}. {match.score * 100:.1f}% — {match.source_doc} — "
            f"{get_match_location_label(match)}"
        )

    return lines
